# Importazione delle librerie necessarie
import os
import logging
from io import BytesIO
import tempfile
import asyncio
import re
import zipfile
import PyPDF2
from docx import Document
from dotenv import load_dotenv
import edge_tts
# Importazioni specifiche per la gestione dell'API OpenAI e del modello
from langchain_openai import ChatOpenAI
from langchain_core.prompts import ChatPromptTemplate

import streamlit as st

# Configurazione del logging per tracciare le operazioni durante l'esecuzione
import logging

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger()

# Messaggio di test per verificare il funzionamento del logging
logger.info("Logging configurato correttamente.")

# Caricamento delle variabili d'ambiente
load_dotenv()

# Recupera la chiave API da .env
openai_api_key = os.getenv("OPENAI_API_KEY")

# Dizionario delle lingue e delle voci disponibili per la sintesi vocale
language_dict = {
    "Italian": {
        "Isabella": "it-IT-IsabellaNeural",
        "Diego": "it-IT-DiegoNeural",
        "Elsa": "it-IT-ElsaNeural"
    },
    "English": {
        "Jenny": "en-US-JennyNeural",
        "Guy": "en-US-GuyNeural",
        "Aria": "en-US-AriaNeural"
    },
    "Swedish": {
        "Sofie": "sv-SE-SofieNeural",
        "Mattias": "sv-SE-MattiasNeural"
    }
}

# Lista dei modelli LLM disponibili
llm_models = [
    "gpt-4o-mini",  # Modello leggero per riassunti rapidi
    "gpt-4o",       # Modello avanzato per riassunti dettagliati
    "gpt-4"         # Modello per task complessi
]

# Impostazioni configurabili per il modello e la sintesi vocale
DEFAULT_LANGUAGE = "Italian"       # Lingua di default per il riassunto e la sintesi vocale
DEFAULT_SPEAKER = "Isabella"       # Voce di default per la sintesi vocale in italiano
DEFAULT_MODEL = llm_models[0]      # Modello LLM di default per il riassunto
DEFAULT_TEMPERATURE = 0.7          # Temperatura di default per il modello

# Funzione per configurare e restituire l'istanza del modello LLM
def configure_llm():
    """Configura e restituisce l'istanza del modello LLM con le impostazioni di default."""
    llm = ChatOpenAI(
        model=DEFAULT_MODEL,
        temperature=DEFAULT_TEMPERATURE,
        openai_api_key=openai_api_key  # La chiave API viene caricata dall'ambiente .env
    )
    return llm

# Funzione per estrarre il testo dal file caricato usando PyPDF2 e python-docx
def load_and_extract_text(uploaded_file):
    try:
        if uploaded_file.type == "text/plain":
            testo = uploaded_file.read().decode("utf-8")
        elif uploaded_file.type == "application/pdf":
            testo = ""
            reader = PyPDF2.PdfReader(uploaded_file)
            for page in reader.pages:
                testo += page.extract_text() + "\n"
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            testo = ""
            doc = Document(uploaded_file)
            for para in doc.paragraphs:
                testo += para.text + "\n"
        else:
            st.error("Formato file non supportato.")
            return None
        return testo
    except Exception as e:
        st.write(f"Errore durante l'estrazione del testo: {e}")
        return None

# Interfaccia Streamlit
st.title("Carica un file (.pdf, .docx, .txt) e visualizza un'anteprima del testo")

# Caricamento del file
uploaded_file = st.file_uploader("Carica il tuo file", type=["pdf", "docx", "txt"])

# Verifica se il file è stato caricato e mostra un'anteprima
if uploaded_file is not None:
    testo_caricato = load_and_extract_text(uploaded_file)

    if testo_caricato:
        anteprima_testo = ' '.join(testo_caricato.split()[:200])
        st.subheader("Anteprima del testo caricato (prime 200 parole):")
        st.write(anteprima_testo)
    else:
        st.write("Nessun testo estratto.")
else:
    st.write("Carica un file per visualizzare il testo.")

# Configurazione dell'LLM globale
llm = configure_llm()

# Funzione per creare un outline utilizzando l'LLM configurato
def create_outline_from_text(text, llm):
    """Genera e restituisce un outline dell'intero testo senza suddivisione in chunk, utilizzando un oggetto LLM preconfigurato."""

    template = ChatPromptTemplate.from_messages([
        ("system", f"""
        Il tuo compito è creare un outline dettagliato del seguente testo in {DEFAULT_LANGUAGE}.

        L'outline deve essere ben strutturato e organizzato gerarchicamente con i seguenti requisiti:
        - Inizia ogni sezione con un titolo chiaro e conciso che riassuma il contenuto della sezione.
        - Se necessario, crea sottosezioni per riflettere la struttura interna del testo.
        - Evita duplicazioni o ripetizioni tra le sezioni e mantieni solo le idee principali.
        - Mantieni l'ordine logico delle idee e la coerenza tra le sezioni.
        - Se il testo contiene esempi o dettagli aggiuntivi, segnalali chiaramente ma in modo conciso.

        Fornisci l'output in forma di elenco strutturato:
        - Usa numeri per le sezioni principali (1, 2, 3, ...)
        - Usa sotto-numeri per le sottosezioni (1.1, 1.2, ...)

        Testo per l'outline:
        """),
        ("human", "{input}")
    ])

    chain = template.pipe(llm)
    response = chain.invoke({
        "input": text
    })

    return response.content

# Generazione dell'outline
st.title("Generazione dell'Outline del Testo Caricato")

outlineA = ""  # Inizializza outlineA come lista vuota

if uploaded_file is not None and testo_caricato:
    with st.spinner("Generazione dell'outline in corso..."):
        outline = create_outline_from_text(testo_caricato, llm)
        st.session_state['outline'] = outline  # Salva outline nello stato della sessione
        st.subheader("Outline Generato:")
        st.write(outline)

        # Aggiungi l'outline generato alla lista outlineA
        outlineA=outline



# Funzione per aggiungere i marker "++++" usando l'LLM
def mark_transitions_with_llm(text, llm):
    """Richiede al modello LLM di inserire '++++' ogni volta che c'è un passaggio tra paragrafi o un cambio di argomento."""

    template = ChatPromptTemplate.from_messages([
        ("system", """
        Analizza il testo e inserisci '++++' ogni volta che c'è un passaggio tra paragrafi o un cambio di argomento.
        Restituisci solo il testo con i marker '++++' nei punti appropriati.
        """),
        ("human", "{input}")
    ])

    chain = template.pipe(llm)
    response = chain.invoke({"input": text})

    return response.content

# Funzione per dividere il testo in chunk usando il separatore "++++"
def split_text_by_marker(text, marker="++++"):
    """Divide il testo in chunk utilizzando il marker specificato."""
    return [chunk.strip() for chunk in text.split(marker) if chunk.strip()]

# Funzione per riassumere ciascun chunk separatamente
def summarize_chunk(chunk, llm):
    """Genera un riassunto conciso del chunk."""

    template = ChatPromptTemplate.from_messages([
        ("system", """
        Il tuo compito è riassumere il testo fornito in modo conciso. Segui queste linee guida:

        - Il riassunto deve essere lungo circa un quarto o un terzo del chunk originale.
        - Utilizza uno stile descrittivo, senza riferimenti diretti come "in questo chunk."
        - Rivolgiti a un pubblico di studenti, concentrandoti sui concetti chiave.
        """),
        ("human", "{input}")
    ])

    chain = template.pipe(llm)
    response = chain.invoke({"input": chunk})

    return response.content

# Analisi e riassunto dei chunk di testo
st.title("Analisi e Riassunto dei Chunk di Testo")

if uploaded_file is not None and testo_caricato:
    with st.spinner("Analisi del testo in corso..."):
        # Aggiunge i marker "++++" e divide il testo
        marked_text = mark_transitions_with_llm(testo_caricato, llm)
        semantic_chunks = split_text_by_marker(marked_text)
        st.session_state['semantic_chunks'] = semantic_chunks

    # Visualizza tutti i chunk generati
    if len(semantic_chunks) > 0:
        st.subheader("Chunk Generati:")
        for i, chunk in enumerate(semantic_chunks, start=1):
            st.subheader(f"Chunk {i}")
            st.write(chunk)
            st.write("-" * 40)
    else:
        st.write("Nessun chunk generato.")

    # Riassunto di ciascun chunk
    summaries = []
    with st.spinner("Generazione dei riassunti in corso..."):
        for i, chunk in enumerate(semantic_chunks):
            summary = summarize_chunk(chunk, llm)
            summaries.append(summary)

    st.session_state['summaries'] = summaries

    # Visualizza tutti i riassunti
    st.subheader("Riassunti dei Chunk:")
    for i, summary in enumerate(summaries, start=1):
        st.subheader(f"Riassunto del Chunk {i}")
        st.write(summary)
        st.write("-" * 40)
else:
    st.write("Carica un file per iniziare l'analisi.")

# Funzione per unire i riassunti
def combine_summaries(summaries):
    """Unisce i riassunti dei chunk in un unico testo e restituisce il testo combinato."""
    combined_text = "\n\n".join(summaries)
    return combined_text

# Funzione per la revisione del testo unito
def revise_combined_text(text, llm):
    """Usa l'LLM per rivedere il testo combinato per migliorare coesione e fluidità."""

    template = ChatPromptTemplate.from_messages([
        ("system", f"""
        Il tuo compito è rivedere il testo fornito, mantenendo tutte le informazioni ma migliorandone la coesione e la fluidità in {DEFAULT_LANGUAGE}.

        Requisiti per la revisione:
        - Mantieni intatte le informazioni e la struttura delle sezioni.
        - Riformula le frasi dove necessario per rendere il testo più scorrevole e coeso.
        - Evita di modificare il contenuto o di omettere informazioni.
        - Inserisci dei titoli per ogni chunk
        Testo da rivedere:
        """),
        ("human", "{input}")
    ])

    chain = template.pipe(llm)
    response = chain.invoke({"input": text})

    return response.content

# Inizializzazione di testA a livello più alto
testA = ""  # testA inizializzato come stringa vuota

# Unione e revisione dei riassunti
if 'summaries' in st.session_state:
    # Unisci i riassunti in un unico testo
    testo_combinato = combine_summaries(st.session_state['summaries'])
    st.session_state['combined_text'] = testo_combinato  # Salva testo combinato

    # Revisione del testo combinato
    with st.spinner("Revisione del testo in corso..."):
        testo_revisionato = revise_combined_text(testo_combinato, llm)
        st.session_state['testo_revisionato'] = testo_revisionato  # Salva testo revisionato

    # Visualizza il testo revisionato
    st.subheader("Testo Revisionato per Coesione e Fluidità:")
    st.write(testo_revisionato)

    # Aggiorna testA con il testo revisionato
    testA = testo_revisionato
################

from io import BytesIO
from docx import Document
import streamlit as st

# Inizializzazione di testA a livello più alto
testA = ""

# Funzione per la preparazione dei file una sola volta
def prepare_files():
    if 'outline' in st.session_state and 'outline_txt' not in st.session_state:
        # Prepara il file TXT per l'outline
        st.session_state['outline_txt'] = BytesIO(st.session_state['outline'].encode('utf-8'))
        
        # Prepara il file DOCX per l'outline
        doc_outline = Document()
        doc_outline.add_paragraph(st.session_state['outline'])
        outline_docx = BytesIO()
        doc_outline.save(outline_docx)
        outline_docx.seek(0)
        st.session_state['outline_docx'] = outline_docx

    if 'testo_revisionato' in st.session_state and 'testo_revisionato_txt' not in st.session_state:
        # Prepara il file TXT per il testo revisionato
        st.session_state['testo_revisionato_txt'] = BytesIO(st.session_state['testo_revisionato'].encode('utf-8'))
        
        # Prepara il file DOCX per il testo revisionato
        doc_revisionato = Document()
        doc_revisionato.add_paragraph(st.session_state['testo_revisionato'])
        revisionato_docx = BytesIO()
        doc_revisionato.save(revisionato_docx)
        revisionato_docx.seek(0)
        st.session_state['testo_revisionato_docx'] = revisionato_docx

# Chiamata alla funzione di preparazione dei file (una sola volta)
if 'outline' in st.session_state and 'testo_revisionato' in st.session_state:
    prepare_files()

# Sezione per il download dei file
st.title("Scarica i file generati")

# Controlla la presenza dei file e visualizza i pulsanti di download
if 'outline_txt' in st.session_state and 'outline_docx' in st.session_state:
    st.download_button("Scarica Outline in TXT", data=st.session_state['outline_txt'], file_name="outline.txt")
    st.download_button("Scarica Outline in DOCX", data=st.session_state['outline_docx'], file_name="outline.docx")

if 'testo_revisionato_txt' in st.session_state and 'testo_revisionato_docx' in st.session_state:
    st.download_button("Scarica Riassunto Revisionato in TXT", data=st.session_state['testo_revisionato_txt'], file_name="riassunto_revisionato.txt")
    st.download_button("Scarica Riassunto Revisionato in DOCX", data=st.session_state['testo_revisionato_docx'], file_name="riassunto_revisionato.docx")
