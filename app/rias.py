# Importazione delle librerie necessarie
import os
import logging
from io import BytesIO
import PyPDF2
from docx import Document
from dotenv import load_dotenv
import streamlit as st
from langchain_openai import ChatOpenAI
from langchain_core.prompts import ChatPromptTemplate

# Configurazione del logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger()
logger.info("Logging configurato correttamente.")

# Caricamento delle variabili d'ambiente
load_dotenv()
openai_api_key = os.getenv("OPENAI_API_KEY")

# Impostazioni di base
DEFAULT_LANGUAGE = "Italian"
DEFAULT_MODEL = "gpt-4o-mini"
DEFAULT_TEMPERATURE = 0.7

# Inizializzazione delle variabili di sessione per il caching
if 'outline' not in st.session_state:
    st.session_state['outline'] = None
if 'testo_revisionato' not in st.session_state:
    st.session_state['testo_revisionato'] = None

# Funzione per configurare il modello LLM
def configure_llm():
    return ChatOpenAI(
        model=DEFAULT_MODEL,
        temperature=DEFAULT_TEMPERATURE,
        openai_api_key=openai_api_key
    )

llm = configure_llm()

# Funzione per estrarre il testo dal file caricato
def load_and_extract_text(uploaded_file):
    try:
        if uploaded_file.type == "text/plain":
            return uploaded_file.read().decode("utf-8")
        elif uploaded_file.type == "application/pdf":
            testo = ""
            reader = PyPDF2.PdfReader(uploaded_file)
            for page in reader.pages:
                testo += page.extract_text() + "\n"
            return testo
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            testo = ""
            doc = Document(uploaded_file)
            for para in doc.paragraphs:
                testo += para.text + "\n"
            return testo
        else:
            st.error("Formato file non supportato.")
            return None
    except Exception as e:
        st.write(f"Errore durante l'estrazione del testo: {e}")
        return None

# Funzione per creare un outline utilizzando l'LLM configurato
def create_outline_from_text(text, llm):
    """Genera e restituisce un outline dell'intero testo."""
    template = ChatPromptTemplate.from_messages([
        ("system", f"""
        Il tuo compito è creare un outline dettagliato del seguente testo in {DEFAULT_LANGUAGE}.
        """),
        ("human", "{input}")
    ])
    chain = template.pipe(llm)
    response = chain.invoke({"input": text})
    return response.content

# Funzione per aggiungere i marker "++++" usando l'LLM
def mark_transitions_with_llm(text, llm):
    """Richiede al modello LLM di inserire '++++' ogni volta che c'è un passaggio tra paragrafi o un cambio di argomento."""
    template = ChatPromptTemplate.from_messages([
        ("system", "Analizza il testo e inserisci '++++' ogni volta che c'è un passaggio tra paragrafi o un cambio di argomento."),
        ("human", "{input}")
    ])
    chain = template.pipe(llm)
    response = chain.invoke({"input": text})
    return response.content

# Funzione per dividere il testo in chunk
def split_text_by_marker(text, marker="++++"):
    return [chunk.strip() for chunk in text.split(marker) if chunk.strip()]

# Funzione per riassumere ciascun chunk separatamente
def summarize_chunk(chunk, llm):
    template = ChatPromptTemplate.from_messages([
        ("system", "Riassumi il testo fornito in modo conciso."),
        ("human", "{input}")
    ])
    chain = template.pipe(llm)
    response = chain.invoke({"input": chunk})
    return response.content

# Funzione per unire i riassunti
def combine_summaries(summaries):
    return "\n\n".join(summaries)

# Funzione per la revisione del testo unito
def revise_combined_text(text, llm):
    template = ChatPromptTemplate.from_messages([
        ("system", f"""
        Il tuo compito è rivedere il testo fornito, migliorando la coesione e la fluidità in {DEFAULT_LANGUAGE}.
        """),
        ("human", "{input}")
    ])
    chain = template.pipe(llm)
    response = chain.invoke({"input": text})
    return response.content

# Caricamento del file
st.title("Carica un file (.pdf, .docx, .txt) e visualizza un'anteprima del testo")
uploaded_file = st.file_uploader("Carica il tuo file", type=["pdf", "docx", "txt"])

if uploaded_file is not None:
    testo_caricato = load_and_extract_text(uploaded_file)
    if testo_caricato:
        st.subheader("Anteprima del testo caricato (prime 200 parole):")
        st.write(' '.join(testo_caricato.split()[:200]))

        # Generazione dell'outline se non è già stato generato
        if st.session_state['outline'] is None:
            with st.spinner("Generazione dell'outline in corso..."):
                outline = create_outline_from_text(testo_caricato, llm)
                st.session_state['outline'] = outline

        # Visualizzazione dell'outline
        st.subheader("Outline Generato:")
        st.write(st.session_state['outline'])

        # Generazione del testo revisionato se non è già stato generato
        if st.session_state['testo_revisionato'] is None:
            # Divisione in chunk e riassunto
            marked_text = mark_transitions_with_llm(testo_caricato, llm)
            semantic_chunks = split_text_by_marker(marked_text)
            summaries = [summarize_chunk(chunk, llm) for chunk in semantic_chunks]
            testo_combinato = combine_summaries(summaries)

            # Revisione del testo combinato
            with st.spinner("Revisione del testo in corso..."):
                testo_revisionato = revise_combined_text(testo_combinato, llm)
                st.session_state['testo_revisionato'] = testo_revisionato

        # Visualizzazione del testo revisionato
        st.subheader("Testo Revisionato:")
        st.write(st.session_state['testo_revisionato'])

# Funzione per preparare i file una sola volta per il download
def prepare_files():
    if 'outline_txt' not in st.session_state:
        st.session_state['outline_txt'] = BytesIO(st.session_state['outline'].encode('utf-8'))
        
        # Salva l'outline in DOCX
        doc_outline = Document()
        doc_outline.add_paragraph(st.session_state['outline'])
        outline_docx = BytesIO()
        doc_outline.save(outline_docx)
        outline_docx.seek(0)
        st.session_state['outline_docx'] = outline_docx

    if 'testo_revisionato_txt' not in st.session_state:
        st.session_state['testo_revisionato_txt'] = BytesIO(st.session_state['testo_revisionato'].encode('utf-8'))
        
        # Salva il testo revisionato in DOCX
        doc_revisionato = Document()
        doc_revisionato.add_paragraph(st.session_state['testo_revisionato'])
        revisionato_docx = BytesIO()
        doc_revisionato.save(revisionato_docx)
        revisionato_docx.seek(0)
        st.session_state['testo_revisionato_docx'] = revisionato_docx

# Preparazione dei file per il download una volta sola
if st.session_state['outline'] and st.session_state['testo_revisionato']:
    prepare_files()

# Sezione di download dei file senza rigenerazione
st.title("Scarica i file generati")

if 'outline_txt' in st.session_state and 'outline_docx' in st.session_state:
    st.download_button("Scarica Outline in TXT", data=st.session_state['outline_txt'], file_name="outline.txt")
    st.download_button("Scarica Outline in DOCX", data=st.session_state['outline_docx'], file_name="outline.docx")

if 'testo_revisionato_txt' in st.session_state and 'testo_revisionato_docx' in st.session_state:
    st.download_button("Scarica Riassunto Revisionato in TXT", data=st.session_state['testo_revisionato_txt'], file_name="riassunto_revisionato.txt")
    st.download_button("Scarica Riassunto Revisionato in DOCX", data=st.session_state['testo_revisionato_docx'], file_name="riassunto_revisionato.docx")
